from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "NUVA_Client_System_Document.docx"


BLACK = RGBColor(0, 0, 0)
DARK = RGBColor(31, 31, 31)
GRAY = RGBColor(85, 85, 85)
LIGHT_BORDER = "DADCE0"
TABLE_HEADER_FILL = "F8F9FA"


def set_run_font(run, name="Calibri", size=11, color=DARK, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table, color=LIGHT_BORDER):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
      borders = OxmlElement("w:tblBorders")
      tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
      el = borders.find(qn(f"w:{edge}"))
      if el is None:
        el = OxmlElement(f"w:{edge}")
        borders.append(el)
      el.set(qn("w:val"), "single")
      el.set(qn("w:sz"), "6")
      el.set(qn("w:space"), "0")
      el.set(qn("w:color"), color)


def configure_page(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def paragraph_border_bottom(paragraph, color=LIGHT_BORDER, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def set_doc_language(doc):
    styles = doc.styles
    styles.element.set(qn("w:val"), "en-US")


def build_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.1

    for style_name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, RGBColor(46, 116, 181)),
        ("Heading 2", 13, 12, 6, RGBColor(46, 116, 181)),
        ("Heading 3", 12, 8, 4, RGBColor(31, 77, 120)),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        fmt = style.paragraph_format
        fmt.space_before = Pt(before)
        fmt.space_after = Pt(after)
        fmt.line_spacing = 1.1

    if "Small Note" not in doc.styles:
        style = doc.styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(9.5)
        style.font.color.rgb = GRAY
        style.paragraph_format.space_after = Pt(4)


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("NUVA | Client System Document")
    set_run_font(run, size=9.5, color=GRAY, bold=False)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("Confidential")
    set_run_font(run, size=9.5, color=GRAY, bold=False)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("CLIENT SYSTEM DOCUMENT")
    set_run_font(run, size=12, color=GRAY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("NUVA E-Commerce Platform")
    set_run_font(run, size=24, color=BLACK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("Current Working System, Role-Based Views, and Proposed Feature Scope")
    set_run_font(run, size=13, color=GRAY, bold=False)

    meta = doc.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    meta.columns[0].width = Inches(1.7)
    meta.columns[1].width = Inches(4.8)
    rows = [
        ("Prepared For", "NUVA stakeholders and client review"),
        ("Document Purpose", "Provide a clean overview of the current system, dummy modules, and proposed next-phase functionality."),
        ("Prepared On", "July 29, 2026"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        row.cells[0].width = Inches(1.7)
        row.cells[1].width = Inches(4.8)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        lp = row.cells[0].paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        lrun = lp.add_run(label)
        set_run_font(lrun, size=10.5, bold=True)
        vp = row.cells[1].paragraphs[0]
        vp.paragraph_format.space_after = Pt(0)
        vrun = vp.add_run(value)
        set_run_font(vrun, size=10.5)
    set_table_borders(meta)

    doc.add_paragraph("")
    note = doc.add_paragraph(style="Small Note")
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(12)
    run = note.add_run(
        "This document distinguishes between features that are live and usable today, features that are currently placeholder or dummy, and features recommended for future rollout."
    )
    set_run_font(run, size=9.5, color=GRAY)

    rule = doc.add_paragraph()
    paragraph_border_bottom(rule)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=11)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_feature_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [Inches(1.6), Inches(1.85), Inches(3.05)]
    headers = ["Section", "Status", "Details"]
    for idx, (cell, header, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, TABLE_HEADER_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=10.5, bold=True)
    for section, status, details in rows:
        row = table.add_row()
        values = [section, status, details]
        for cell, value, width in zip(row.cells, values, widths):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=10.25)
    set_table_borders(table)


def add_role_section(doc, title, intro, current_features, item_views, dummy_features, proposed_features):
    doc.add_heading(title, level=1)
    add_body_paragraph(doc, intro)

    doc.add_heading("Current Features", level=2)
    add_bullets(doc, current_features)

    doc.add_heading("Item Views", level=2)
    add_bullets(doc, item_views)

    doc.add_heading("Dummy or Placeholder Features", level=2)
    add_bullets(doc, dummy_features)

    doc.add_heading("Proposed Features", level=2)
    add_bullets(doc, proposed_features)


def build_document():
    doc = Document()
    configure_page(doc.sections[0])
    build_styles(doc)
    add_header_footer(doc.sections[0])
    add_title_page(doc)

    doc.add_heading("1. Executive Overview", level=1)
    add_body_paragraph(
        doc,
        "NUVA is a full-stack jewelry e-commerce platform with three main operating views: the public website visitor experience, the logged-in customer experience, and the internal admin or super admin experience. The current system already supports storefront browsing, product management, cart and checkout flow, and core admin catalog operations."
    )
    add_body_paragraph(
        doc,
        "This document organizes the platform by user role and by feature division so stakeholders can quickly see what is already working, what is currently a placeholder, and what can be prioritized in the next development phase."
    )

    doc.add_heading("2. Platform Structure", level=1)
    add_bullets(
        doc,
        [
            "Frontend: React + Vite + Ant Design",
            "Backend: FastAPI",
            "Database: MongoDB Atlas",
            "Media Storage: Backblaze B2 with Cloudflare CDN delivery",
            "Primary Views: Visitor, Logged-In Customer, Admin / Super Admin",
        ],
    )

    add_role_section(
        doc,
        "3. Website Visitor View",
        "This is the public-facing experience for anyone visiting the NUVA storefront without logging in.",
        [
            "Homepage with brand-led hero section and featured products",
            "Shop page with searchable product listing",
            "Category-based product filtering",
            "Product detail pages with pricing, stock, and material information",
            "Add to cart from listing pages and product pages",
            "Header cart drawer and full cart page",
            "Login and registration entry points",
        ],
        [
            "Product listing card view",
            "Product detail page view",
            "Cart drawer item view",
            "Full cart page item view",
        ],
        [
            "Product fallback data may appear when API requests fail",
            "Cart drawer formatting is functional but less consistent than the full cart page",
        ],
        [
            "Wishlist or save-for-later capability",
            "Advanced filtering by price, material, stone, and collection",
            "Related products and recently viewed products",
            "Product review and rating system",
            "Enhanced product image gallery with zoom and multiple images",
            "Guest checkout option",
        ],
    )

    add_role_section(
        doc,
        "4. Logged-In Customer View",
        "This view is available to registered users after login and supports the purchase and order lifecycle.",
        [
            "Persistent authenticated session",
            "Protected checkout access",
            "Checkout form with customer details, address, and payment method selection",
            "Order placement and cart clearing after successful checkout",
            "My Orders page with order history and status visibility",
            "Logout capability",
        ],
        [
            "Checkout order summary view",
            "Order history table view",
            "Cart items with quantity controls",
            "Purchased order records",
        ],
        [
            "Payment options are currently simple workflow placeholders rather than a full payment gateway integration",
            "No dedicated customer profile dashboard yet",
            "No saved addresses, saved cards, returns, reorder, or invoice tools yet",
        ],
        [
            "Customer account dashboard",
            "Saved addresses and saved payment methods",
            "Order tracking, cancellation, and return requests",
            "Invoice download",
            "Reorder previous purchases",
            "Loyalty or rewards program",
        ],
    )

    doc.add_heading("5. Admin / Super Admin View", level=1)
    add_body_paragraph(
        doc,
        "The admin environment is divided into four major operational areas: Storefront, Commerce, Operations, and Account. Several modules are already functional today, especially in product, category, inventory, and order administration."
    )

    doc.add_heading("5.1 Admin Divisions Summary", level=2)
    add_feature_table(
        doc,
        [
            ("Storefront", "Partly Live", "Dashboard UI, orders, products, categories, and inventory are active. Customers is still placeholder."),
            ("Commerce", "Partly Live", "Orders, products, categories, and inventory are active. Discounts, reviews, payments, returns, and customers remain placeholder."),
            ("Operations", "Mostly Placeholder", "Dashboard shell exists. Admins or staff, website content, media library, and shipping are not yet implemented."),
            ("Account", "Partly Live", "Logout works. Settings, reports, and profile are still placeholder or dummy."),
        ],
    )

    doc.add_heading("5.2 Storefront Division", level=2)
    add_bullets(
        doc,
        [
            "Subdivisions: Dashboard, Orders, Products, Categories, Inventory, Customers",
            "Working now: order management, product catalog management, category management, and inventory updates",
            "Product module supports add, edit, duplicate, visibility toggle, archive, soft delete, and preview",
            "Category module supports add, edit, delete, activate or deactivate, and sorting",
            "Inventory module supports stock quantity and low-stock threshold updates",
            "Customers page is currently placeholder",
            "Dashboard content is largely static summary content at present",
        ],
    )

    doc.add_heading("5.3 Commerce Division", level=2)
    add_bullets(
        doc,
        [
            "Subdivisions: Dashboard, Orders, Products, Categories, Inventory, Customers, Discounts, Reviews, Payments, Returns",
            "Working now: orders, products, categories, and inventory",
            "Placeholder today: customers, discounts, reviews, payments, and returns",
            "Recommended next build: coupons, review moderation, refund handling, payment logs, and returns workflow",
        ],
    )

    doc.add_heading("5.4 Operations Division", level=2)
    add_bullets(
        doc,
        [
            "Subdivisions: Dashboard, Admins or Staff, Website Content, Media Library, Shipping",
            "Current state: dashboard shell exists, but the operational modules themselves are not yet built",
            "Recommended next build: staff roles and permissions, content management tools, media library, and shipping configuration",
        ],
    )

    doc.add_heading("5.5 Account Division", level=2)
    add_bullets(
        doc,
        [
            "Subdivisions: Dashboard, Reports, Settings, Profile, Logout",
            "Working now: admin logout and top-right admin dropdown access",
            "Settings entry is currently a dummy action",
            "Reports and profile pages are still placeholder screens",
            "Recommended next build: business settings, profile management, analytics reports, and security controls",
        ],
    )

    doc.add_heading("6. Item and Shopping Flow Views", level=1)
    add_body_paragraph(
        doc,
        "The platform already includes multiple item-based views that support browsing, purchase, and administration."
    )
    add_bullets(
        doc,
        [
            "Public product listing cards",
            "Public product detail page",
            "Header cart drawer items",
            "Full shopping cart page items",
            "Checkout summary item list",
            "Admin product table rows",
            "Admin product editor or drawer",
            "Inventory table rows",
            "Draft image tiles for incomplete products",
        ],
    )
    add_body_paragraph(
        doc,
        "Recommended future additions include quick-view product modals, product comparison, variant selection, related item sections, and admin bulk-edit item views."
    )

    doc.add_heading("7. Role Summary", level=1)
    add_feature_table(
        doc,
        [
            ("Website Visitor", "Live", "Can browse products, search, filter, view product pages, and add items to cart."),
            ("Logged-In Customer", "Live", "Can checkout, place orders, view order history, and log out."),
            ("Admin / Super Admin", "Live + Partial", "Can manage products, categories, inventory, and orders. Several advanced modules remain placeholder."),
        ],
    )

    doc.add_heading("8. Recommended Next Phase Priorities", level=1)
    add_numbered(
        doc,
        [
            "Customer account dashboard and profile tools",
            "Admin customer management module",
            "Discounts and coupon system",
            "Reviews and moderation workflow",
            "Shipping settings and delivery rules",
            "Reports and analytics dashboards",
            "Account settings and profile completion",
            "Website content and media management",
        ],
    )

    doc.add_heading("9. Conclusion", level=1)
    add_body_paragraph(
        doc,
        "NUVA already has a solid operational base for storefront browsing, cart and checkout flow, order placement, and core admin catalog management. The strongest completed area today is the admin product and catalog system, supported by category and inventory tools."
    )
    add_body_paragraph(
        doc,
        "The next phase should focus on completing the placeholder operational modules, expanding customer account functionality, and strengthening finance, shipping, content, and reporting capabilities so the platform can operate as a fully rounded production commerce system."
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
